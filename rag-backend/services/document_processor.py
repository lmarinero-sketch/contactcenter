"""
Document Processor — Extracts text from various file formats
Includes OCR fallback for scanned PDFs using Tesseract
"""
import os
import re
import csv
import io


def extract_text(file_path: str) -> str:
    """Extract text from a file based on its extension."""
    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
        '.pdf': _extract_pdf,
        '.docx': _extract_docx,
        '.xlsx': _extract_excel,
        '.xls': _extract_excel,
        '.csv': _extract_csv,
    }

    extractor = extractors.get(ext, _extract_text_file)
    raw_text = extractor(file_path)
    return clean_text(raw_text)


def _is_meaningful_text(text: str) -> bool:
    """Check if extracted text has real content (not just DocuSign IDs, headers, etc.)"""
    if not text:
        return False
    # Filter out lines that are only DocuSign envelope IDs or whitespace
    meaningful_lines = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Skip DocuSign envelope ID lines
        if 'DocuSign Envelope ID' in line:
            continue
        # Skip very short lines (page numbers, etc.)
        if len(line) < 5:
            continue
        meaningful_lines.append(line)
    return len('\n'.join(meaningful_lines)) > 100


def _extract_pdf(file_path: str) -> str:
    """
    Extract text from PDF using multiple strategies:
    1. pdfplumber (best for complex/signed PDFs like DocuSign)
    2. PyPDF2 (fallback for simpler PDFs)
    3. OCR via Tesseract (last resort for scanned PDFs)
    """
    # Strategy 1: pdfplumber (handles DocuSign, complex layouts)
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    # Filter DocuSign envelope IDs
                    clean_lines = [l for l in text.split('\n')
                                   if 'DocuSign Envelope ID' not in l]
                    clean_text = '\n'.join(clean_lines).strip()
                    if clean_text:
                        pages.append(f"[P\u00e1gina {i}]\n{clean_text}")
        extracted = "\n\n".join(pages)
        if _is_meaningful_text(extracted):
            print(f"PDF: Extracted {len(extracted)} chars via pdfplumber")
            return extracted
    except ImportError:
        print("PDF: pdfplumber not installed, trying PyPDF2...")
    except Exception as e:
        print(f"PDF: pdfplumber failed: {e}, trying PyPDF2...")

    # Strategy 2: PyPDF2 (fallback)
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                clean_lines = [l for l in text.split('\n')
                               if 'DocuSign Envelope ID' not in l]
                clean_text = '\n'.join(clean_lines).strip()
                if clean_text:
                    pages.append(f"[P\u00e1gina {i}]\n{clean_text}")
        extracted = "\n\n".join(pages)
        if _is_meaningful_text(extracted):
            print(f"PDF: Extracted {len(extracted)} chars via PyPDF2")
            return extracted
    except Exception as e:
        print(f"PDF: PyPDF2 failed: {e}")

    # Strategy 3: OpenAI Vision for scanned PDFs (more accurate than Tesseract)
    print(f"PDF: No meaningful text found, trying OpenAI Vision...")
    vision_result = _extract_pdf_vision(file_path)
    if _is_meaningful_text(vision_result):
        return vision_result

    # Strategy 4: Tesseract OCR (last resort, free but less accurate)
    print(f"PDF: Vision failed or unavailable, trying Tesseract OCR...")
    return _extract_pdf_ocr(file_path)


def _extract_pdf_vision(file_path: str, max_pages: int = 10) -> str:
    """
    Extract text from PDF pages using GPT-4o Vision.
    Converts pages to images and sends them to OpenAI's vision API.
    Limited to max_pages to control costs.
    """
    try:
        import base64
        from pdf2image import convert_from_path
        from config import openai_client, CHAT_MODEL

        images = convert_from_path(file_path, dpi=200)
        print(f"Vision: Processing {min(len(images), max_pages)}/{len(images)} pages...")

        pages = []
        for i, image in enumerate(images[:max_pages], 1):
            # Convert PIL image to base64
            import io as _io
            buffer = _io.BytesIO()
            image.save(buffer, format="PNG")
            img_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")

            try:
                response = openai_client.chat.completions.create(
                    model=CHAT_MODEL,
                    max_tokens=2000,
                    timeout=60,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": (
                                        "Extraé TODO el texto de esta página de documento. "
                                        "Mantené la estructura original (títulos, párrafos, listas, tablas). "
                                        "Respondé SOLO con el texto extraído, sin comentarios."
                                    )
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{img_base64}",
                                        "detail": "high"
                                    }
                                }
                            ]
                        }
                    ]
                )
                text = response.choices[0].message.content.strip()
                if text:
                    pages.append(f"[Página {i}]\n{text}")
                print(f"Vision: Page {i} processed ({len(text)} chars)")
            except Exception as e:
                print(f"Vision: Page {i} failed: {e}")

        result = "\n\n".join(pages)
        print(f"Vision: Total extracted: {len(result)} chars from {len(pages)} pages")
        return result

    except ImportError as e:
        print(f"Vision: Missing dependency: {e}")
        return ""
    except Exception as e:
        print(f"Vision extraction failed: {e}")
        return ""


def _extract_pdf_ocr(file_path: str) -> str:
    """Extract text from scanned PDF using Tesseract OCR."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        # Convert PDF pages to images
        images = convert_from_path(file_path, dpi=300)
        print(f"OCR: Converting {len(images)} pages...")

        pages = []
        for i, image in enumerate(images, 1):
            # Run Tesseract OCR with Spanish + English
            text = pytesseract.image_to_string(image, lang='spa+eng')
            if text and text.strip():
                pages.append(f"[Página {i} — OCR]\n{text.strip()}")
            print(f"OCR: Page {i}/{len(images)} processed ({len(text)} chars)")

        result = "\n\n".join(pages)
        print(f"OCR: Total extracted: {len(result)} chars from {len(pages)} pages")
        return result

    except Exception as e:
        print(f"OCR failed: {e}")
        return ""


def _extract_docx(file_path: str) -> str:
    """Extract text from Word documents."""
    from docx import Document

    doc = Document(file_path)
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _extract_excel(file_path: str) -> str:
    """
    Extract text from Excel files with semantic context.
    Each row includes column headers as key-value pairs so the LLM
    understands what each cell means (critical for RAG retrieval).
    """
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"\n=== Hoja: {sheet_name} ===\n")

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # Detect headers: first non-empty row
        headers = None
        data_start = 0
        for i, row in enumerate(rows):
            cells = [str(c).strip() if c is not None else "" for c in row]
            non_empty = [c for c in cells if c]
            if len(non_empty) >= 2:  # At least 2 non-empty cells = likely header
                headers = cells
                data_start = i + 1
                break

        if headers:
            # Output header line for context
            header_line = " | ".join([h for h in headers if h])
            parts.append(f"Columnas: {header_line}\n")

            # Output each data row as key-value pairs for semantic clarity
            for row in rows[data_start:]:
                cells = [str(c).strip() if c is not None else "" for c in row]
                # Skip completely empty rows
                non_empty = [c for c in cells if c]
                if not non_empty:
                    continue

                # Build key-value representation
                kv_parts = []
                for j, cell_val in enumerate(cells):
                    if cell_val and j < len(headers) and headers[j]:
                        kv_parts.append(f"{headers[j]}: {cell_val}")
                    elif cell_val:
                        kv_parts.append(cell_val)

                if kv_parts:
                    parts.append(" | ".join(kv_parts))
        else:
            # No headers detected, fallback to plain extraction
            for row in rows:
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_csv(file_path: str) -> str:
    """Extract text from CSV files."""
    parts = []
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        for row in reader:
            cells = [c.strip() for c in row if c.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text_file(file_path: str) -> str:
    """Extract text from plain text files (.txt, .md, .json, .xml, .html)."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def clean_text(text: str) -> str:
    """Sanitize text: remove invalid Unicode, null bytes, control characters."""
    if not text:
        return ""

    # Remove invalid Unicode surrogates
    text = text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')

    # Remove null bytes
    text = text.replace('\x00', '')

    # Remove control characters but preserve \n and \t
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    return text.strip()
